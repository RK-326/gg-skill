# -*- coding: utf-8 -*-
import re, os, json, glob
from struct import unpack
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/renatakoviazina/Documents/Global Generation/ПДФ на портале"
RAW = open('/tmp/bak_raw.txt', encoding='utf-8').read()
FC = json.load(open(BASE + "/research/factcheck_perplexity.json", encoding="utf-8"))
DATE = "9 июня 2026"
def src(cid): return "; ".join(FC.get(cid, {}).get("sources", []))

HEAD_RE = re.compile(r'^\s*\d{1,2}\s*\.\s*(ВЕСНА|ИЮЛЬ|АВГУСТ|СЕНТЯБР|ОКТЯБР|НОЯБР|ДЕКАБР|ЯНВАР|ФЕВРАЛ|МАРТ)')
pages = RAW.split('\x0c')
pageblk = []
for pi, page in enumerate(pages):
    kept = [ln for ln in page.split('\n') if not re.fullmatch(r'\d{1,2}', ln.strip())]
    for b in re.split(r'\n\s*\n', '\n'.join(kept)):
        b = ' '.join(x.strip() for x in b.split('\n') if x.strip())
        b = re.sub(r'[​‌‍‎‏]', '', b)
        b = re.sub(r'\s+', ' ', b).strip()
        if b: pageblk.append((b, pi+1))
blocks = []
for k,(b,pg) in enumerate(pageblk):
    if blocks and pageblk[k-1][1] != pg:
        prev = blocks[-1][0]
        cont = re.match(r'^[a-zа-яё]', b) or re.match(r'^-\s*\d', b) or re.match(r'^\d', b)
        if cont and not re.search(r'[.!?:»)\]"]\s*$', prev) and not HEAD_RE.match(b):
            blocks[-1] = (prev + ' ' + b, blocks[-1][1]); continue
    blocks.append((b, pg))

STRAY = {"качества","характера","(упорство,","трудолюбие,","стойкость,"}
fixed = []
for b,pg in blocks:
    if b.strip() in STRAY: continue
    if "свои главные креативность и т. д.)." in b:
        b = b.replace("свои главные креативность",
                      "свои главные качества характера (упорство, трудолюбие, стойкость, креативность")
    fixed.append((b,pg))
blocks = fixed

HEADINGS = [
 (1,"ВЕСНА-ЛЕТО","Стажировки, олимпиады, сотрудничество с профессорами"),
 (2,"ИЮЛЬ-АВГУСТ (шаг 1)","Поиск университетов, финансовая помощь"),
 (3,"ИЮЛЬ-АВГУСТ (шаг 2)","Как составить итоговый список университетов"),
 (4,"ИЮЛЬ-АВГУСТ (шаг 3)","Все о стипендиях и финансовой помощи. Как учиться бесплатно?"),
 (5,"АВГУСТ","Исследуем сайты университетов"),
 (6,"АВГУСТ-СЕНТЯБРЬ (шаг 1)","Заполняем Honors and Activities. Какие достижения и активности нужны?"),
 (7,"АВГУСТ-СЕНТЯБРЬ (шаг 2)","Регистрируемся на основном портале для поступления «CommonApp». Получаем fee waivers"),
 (8,"АВГУСТ-СЕНТЯБРЬ (шаг 3)","Пишем главное эссе"),
 (9,"СЕНТЯБРЬ","Пишем эссе «Why this university?», разрабатываем универсальный шаблон для всех университетов"),
 (10,"СЕНТЯБРЬ","Составляем портфолио"),
 (11,"СЕНТЯБРЬ-ОКТЯБРЬ","Все о SAT: подготовка, все полезные материалы, как, где и когда сдавать"),
 (12,"НОЯБРЬ","Проверка первых двух видов эссе. Пишем эссе Supplementals"),
 (13,"НОЯБРЬ-ДЕКАБРЬ","Пишем эссе Extracurriculars"),
 (14,"ДЕКАБРЬ","Собираем документы"),
 (15,"ДЕКАБРЬ","Отправляем результаты SAT"),
 (16,"ДЕКАБРЬ","Тест на знание английского (IELTS/TOEFL/DET) и все ресурсы для бесплатной подготовки"),
 (17,"ДЕКАБРЬ-ЯНВАРЬ","Заполняем CSS Профиль"),
 (18,"ДЕКАБРЬ-МАРТ","Подаем заявки и ждем новостей!"),
]
RAW_HEAD = [(n, re.sub(r'\s*\(шаг \d\)','',m), t) for (n,m,t) in HEADINGS]
HDISP = {n: f"{n}. {m}. {t}" for (n,m,t) in HEADINGS}

elements = []
for b,pg in blocks:
    hit = False
    for (num,mk,title) in RAW_HEAD:
        m = re.match(r'^\s*%d\s*\.\s*%s\s*\.\s*%s' % (num, re.escape(mk), re.escape(title)), b)
        if m:
            elements.append(('h', num, pg))
            rest = b[m.end():].strip()
            if rest: elements.append(('p', rest, pg))
            hit = True; break
    if not hit: elements.append(('p', b, pg))
assert sorted(n for t,n,_ in elements if t=='h') == list(range(1,19)), "headings!"

EDITS = [
 ("Здесь можно приукрасить свои достижения, но прям совсем чуть-чуть. Обычно их можно сделать более значительными, если просто хорошо описать без лжи. Но если очень хочется, можно приврать, например, о конкуренции (то есть описывая свои победы в олимпиадах, можно написать, что конкурентов (участников олимпиады) было около десяти тысяч, а вы оказались одним из нескольких победителей. Тут никто точных чисел не знает). Но полностью выдумывать достижения или деятельность не стоит, так как скорее всего, если вы их не сможете доказать ни дипломами, ни какими-либо ссылками, ни рекомендательными письмами, это просто не будут рассматривать серьезно.",
  "Здесь можно подать свои достижения выигрышнее, но без вранья. Обычно достаточно просто хорошо описать то, что реально было, и достижение уже звучит весомее. Выдумывать достижения или деятельность не стоит: если вы не сможете подтвердить их дипломами, ссылками или рекомендательными письмами, заявку просто не будут рассматривать всерьёз."),
 ("Важно, чтобы эссе было историей из жизни (ее, конечно, можно выдумать, или же можно приукрасить реальную историю, главное, чтобы было правдоподобно).",
  "Важно, чтобы эссе было историей из жизни. Реальную историю можно подать ярче и выразительнее, главное, чтобы звучало правдоподобно."),
 ("На всякий случай: если у вас не было никаких ситуаций в жизни, о которых можно рассказать в подобных эссе, вы спокойно можете все выдумать.",
  "Если яркой ситуации сходу не вспоминается, подумайте о своих привычках, увлечениях и обычных днях: почти всегда находится что-то живое и личное, о чём можно рассказать."),
 (" Тут можно и приврать :)", ""),
 ("Вообще, почти все университеты сейчас test-optional, то есть вы не обязательно должны отправлять им свои результаты SAT.",
  "Сейчас многие университеты test-optional, то есть отправлять результаты SAT не обязательно. Но это быстро меняется: часть вузов уже вернула обязательный SAT, поэтому требования каждого вуза проверяйте на его сайте."),
 ("то в переводе на американскую систему будет: 1) 5 : 4 = 1.25 (коэффициент, на который надо будет делить ваш средний балл); 2) 4.5 : 1.25 = 3.6 (то есть, чтобы посчитать свой средний балл по американской системе, надо взять свой средний балл по пятибалльной системе и разделить его на 1.25). То есть, максимальный GPA - 4.0.",
  "перевести его можно так. Сначала находим коэффициент: 5 : 4 = 1.25. Потом делим на него свой балл: 4.5 : 1.25 = 3.6. Максимальный GPA по этой системе - 4.0."),
 ("Need-blind университеты: ●",
  "Важная оговорка: многие вузы называют себя need-blind, но имеют в виду только граждан США, а финансы иностранных абитуриентов всё-таки учитывают. Need-blind именно для интернациональных студентов, это короткий список ниже. Все остальные вузы на этапе планирования считайте need-aware, пока на их сайте не написано обратное. Список периодически меняется, статус конкретного вуза лучше перепроверять на его официальной странице. Need-blind университеты: ●"),
 ("Brown University (начиная с 2025 года)", "Brown University (для поступающих с набора 2025 года, проверьте на сайте)"),
 ("University of Notre Dame (начиная с 2025 года)", "University of Notre Dame"),
 ("кроме России и Беларуси.", "кроме России и Беларуси (правила и доступные страны меняются, уточните на collegeboard.org)."),
 ("регистрировались на SAT (https://www.collegeboard.org)", "регистрировались (https://www.collegeboard.org)"),
]
def apply_edit(old,new):
    for i,(t,txt,pg) in enumerate(elements):
        if t=='p' and old in txt:
            elements[i]=('p', txt.replace(old,new), pg); return True
    print("EDIT-MISS:", old[:50]); return False
for o,n in EDITS: apply_edit(o,n)
for i,(t,txt,pg) in enumerate(elements):
    if t=='p' and 'undetectable.ai/?' in txt:
        elements[i]=('p', re.sub(r'https://undetectable\.ai/\?\S*','https://undetectable.ai/',txt), pg)

SPLIT_BEFORE = ["Также есть такие понятия, как weighted GPA","- минимальный требуемый результат SAT",
 "- acceptance rate","Примеры merit-based aid:","Need-based aid дают исходя","Примеры need-based aid:",
 "Эти университеты имеют значительные финансовые ресурсы","Need-aware университеты: финансовое положение",
 "Большинство университетов в США являются need-aware","Need-blind университеты - отличный выбор",
 "Помимо этого, про каждый университет"]
def split_block(text):
    segs=[text]
    for m in SPLIT_BEFORE:
        out=[]
        for s in segs: out.extend(re.split(r'(?=%s)'%re.escape(m), s))
        segs=out
    return [s.strip() for s in segs if s.strip()]
_new=[]
for t,v,pg in elements:
    if t=='p':
        for piece in split_block(v): _new.append(('p',piece,pg))
    else: _new.append((t,v,pg))
elements=_new

def insert_after(anchor, typ, payload):
    for i,(t,txt,pg) in enumerate(elements):
        if t=='p' and anchor in txt:
            elements.insert(i+1, (typ, payload, pg)); return True
    print("INS-MISS:", typ, "::", anchor[:42]); return False

for i,(t,v,pg) in enumerate(elements):
    if t=='h':
        elements.insert(i, ('genpic', '/tmp/bak_imgs/illus_1.png', pg)); break
ILLUS = [
 ("Bachelor of Nursing", 'table', 'degrees'),
 ("Максимальный GPA по этой системе", 'genpic', '/tmp/bak_imgs/illus_3.png'),
 ("Need-based aid предоставляется на основе финансовых", 'table', 'meritneed'),
 ("Washington and Lee University", 'table', 'needblind'),
 ("3-5 safe choices", 'genpic', '/tmp/bak_imgs/illus_6.png'),
 ("Здесь можно подать свои достижения выигрышнее", 'table', 'honors'),
 ("Не бойтесь экспериментировать в эссе", 'genpic', '/tmp/bak_imgs/illus_8.png'),
 ("лучше всего 1500+", 'genpic', '/tmp/bak_imgs/illus_9.png'),
 ("Теперь собираем документы! Что нужно собрать", 'table', 'docs'),
 ("Regular Decision (Rolling Admission)", 'table', 'eaedrd'),
 ("1500 - 3000 долларов", 'genpic', '/tmp/bak_imgs/illus_12.png'),
]
for a,k,p in ILLUS: insert_after(a,k,p)

FN = [
 ("Максимальный GPA по этой системе","gpa","деление среднего балла на 1.25 это лишь грубый ориентир «для себя»; официально оценки переводят по каждому предмету (A=4, B=3), часто через сервис оценки документов (WES и аналоги)."),
 ("за последние четыре года","gpa_years","GPA для бакалавриата обычно считают по последним годам школы; точный набор лет и предметов зависит от вуза."),
 ("часть вузов уже вернула обязательный SAT","test_optional","«почти все test-optional» уже неверно: ряд вузов вернул обязательный SAT/ACT на 2025-2026."),
 ("принимают, как правило, меньше 10%","acceptance_rate","у самых селективных вузов acceptance rate ниже 10%; «типичные» 30% сильно зависят от вуза."),
 ("будет называться major","major_minor","major это основная специальность, minor это укороченная дополнительная программа сверх major."),
 ("BN (Bachelor of Nursing)","degree_types","расшифровки BA/BFA/BS/BAS/BBA/BEng/BArch/BN подтверждены."),
 ("автоматически рассматривают всех абитуриентов как кандидатов","merit_auto","на merit-aid обычно рассматривают автоматически; часть стипендий требует отдельной заявки и эссе."),
 ("Pell Grant (недоступны без гражданства США или грин карты)","pell","федеральный Pell Grant доступен только гражданам США и отдельным категориям неграждан (например, грин-карта); иностранным студентам нет."),
 ("Washington and Lee University","need_blind","для иностранных студентов need-blind лишь у короткого списка вузов; список и статусы меняются, проверяйте на сайте вуза."),
 ("Need-aware университеты могут быть более конкурентными","need_aware_def","need-blind: финансы не влияют на приём; need-aware: учитываются, что снижает шансы нуждающихся."),
 ("от 35 000 до 100 000 долларов","tuition_cost","годовая стоимость обучения и проживания без помощи попадает в диапазон примерно 35-100 тыс. долларов."),
 ("Это все нужно будет занести в вашу анкету в CommonApp","commonapp_activities","в Common App до 10 пунктов Activities и до 5 Honors, с лимитом символов на каждое."),
 ("А теперь начинаем писать main essay","main_essay_words","главное эссе Common App до 650 слов."),
 ("Не бойтесь экспериментировать в эссе","supplemental_essays","supplemental и «Why this college» эссе задаёт каждый вуз отдельно, лимиты слов различаются."),
 ("лучше всего 1500+","sat_structure","цифровой SAT: два модуля Reading and Writing и два модуля Math, второй модуль адаптивный, максимум 1600."),
 ("Заполняете CSS Профиль на уже знакомом сайте","css_profile","CSS Profile на новый цикл открывается 1 октября."),
 ("ОГРОМНАЯ анкета о вашем финансовом положении","css_cost","подача CSS Profile платная: стартовый взнос и плата за каждый дополнительный вуз (для части семей есть waiver)."),
 ("Все это отнесите в нотариальный перевод","rec_letters","вузам обычно нужны рекомендации от counselor и одного-двух учителей; точное число зависит от вуза."),
 ("кроме России и Беларуси (правила и доступные страны","sat_ru_by","центров SAT в России и Беларуси нет, нужно ехать в другую страну."),
 ("ОЧЕНЬ советую сдавать Дуолинго Тест","english_tests_accept","TOEFL, IELTS и Duolingo English Test принимает большинство вузов; конкретный список проверяйте на сайте."),
 ("Regular Decision (Rolling Admission)","ed_binding","Early Decision обязывает поступить и подаётся только в один вуз; Early Action не обязывает."),
 ("следите за дедлайнами подачи заявок","app_deadline","типичный дедлайн Regular Decision около 1 января; у каждого вуза свой."),
 ("application fee и могли подать заявку через","fee_waiver","fee waiver в Common App доступен при финансовой нужде по самоподтверждению или через counselor."),
 ("1500 - 3000 долларов","costs","суммы ориентировочные: цены тестов и application fees меняются; проверяйте на официальных сайтах."),
 ("https://www.efset.org/ru/","efset_free","EF SET (efset.org) это бесплатный онлайн-тест уровня английского."),
 ("если вы из России или Беларуси","grammarly_ru","Grammarly ограничивает доступ для пользователей из России и Беларуси из-за санкций."),
 ("отправляете свои результаты SAT во все университеты","sat_superscore","многие вузы используют superscore SAT, складывая лучшие результаты секций за разные даты."),
 ("стоить от 35 000 до 100 000 долларов","toefl_ielts_cost","TOEFL iBT и IELTS стоят порядка 200-350 долларов; Duolingo заметно дешевле."),
]
for anchor,cid,phrase in FN:
    if not src(cid): print("NO-SRC:",cid); continue
    text = re.sub(r'\s+',' ', f"Проверено (Perplexity, {DATE}): {phrase} Источники: {src(cid)}")
    insert_after(anchor,'fn',text)

# embed real screenshots by page
def jpeg_size(path):
    with open(path,'rb') as fh:
        fh.read(2)
        while True:
            b=fh.read(1)
            while b and b!=b'\xff': b=fh.read(1)
            marker=fh.read(1)
            if marker in (b'\xc0',b'\xc1',b'\xc2',b'\xc3'):
                fh.read(3); h,w=unpack('>HH',fh.read(4)); return w,h
            elif marker in (b'\xd9',b''): return 0,0
            else:
                ln=unpack('>H',fh.read(2))[0]; fh.read(ln-2)
real={}
for f in sorted(glob.glob('/tmp/bak_imgs/img-*.jpg')):
    p=int(re.match(r'.*img-(\d+)-(\d+)\.jpg',f).group(1))
    w,h=jpeg_size(f)
    if (w,h) in ((1545,2000),(2048,485)) or w==0: continue
    real.setdefault(p,[]).append((f,w,h))
for p in sorted(real):
    idx=None
    for i,(t,v,pg) in enumerate(elements):
        if pg==p: idx=i
    if idx is None:
        for i,(t,v,pg) in enumerate(elements):
            if pg and pg<=p: idx=i
    if idx is None: idx=len(elements)-1
    for tup in real[p]:
        idx+=1; elements.insert(idx, ('pic',tup,p))

def clean(s): return re.sub(r'\s+',' ', s.replace('—','-').replace('–','-')).strip()
doc=Document()
doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(11.5)
def shade(par,color):
    pPr=par._p.get_or_add_pPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color); pPr.append(sh)
def shade_cell(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),color); tcPr.append(sh)
def make_table(header, rows):
    cap=doc.add_paragraph('(иллюстрация, сгенерирована автоматически)')
    cap.runs[0].italic=True; cap.runs[0].font.size=Pt(8); cap.runs[0].font.color.rgb=RGBColor(0x88,0x88,0x88)
    t=doc.add_table(rows=1+len(rows), cols=len(header)); t.style='Table Grid'
    for j,h in enumerate(header):
        c=t.rows[0].cells[j]; c.text=''; r=c.paragraphs[0].add_run(clean(h))
        r.bold=True; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r.font.size=Pt(9.5); shade_cell(c,'1565C0')
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            c=t.rows[i+1].cells[j]; c.text=''; r=c.paragraphs[0].add_run(clean(val)); r.font.size=Pt(9)
    doc.add_paragraph()
TABLES={
 'degrees':(['Аббр.','Расшифровка и направление'],[
   ['BA','Bachelor of Arts: гуманитарные и социальные науки, языки'],
   ['BFA','Bachelor of Fine Arts: изобразительное искусство, театр, кино'],
   ['BS','Bachelor of Science: естественные науки, инженерия, математика'],
   ['BAS','Bachelor of Applied Science: прикладные технические программы'],
   ['BBA','Bachelor of Business Administration: бизнес, финансы, менеджмент'],
   ['BEng','Bachelor of Engineering: инженерные науки'],
   ['BArch','Bachelor of Architecture: архитектура'],
   ['BN','Bachelor of Nursing: сестринское дело, медицина']]),
 'meritneed':(['Merit-based (за заслуги)','Need-based (по нуждам семьи)'],[
   ['Учёба: оценки, SAT, олимпиады, проекты','Доход, имущество и обязательства семьи'],
   ['Спорт: выступления нац. и межд. уровня','Стоимость учёбы и жизни в вузе'],
   ['Творчество: награды в искусстве','Финансовые формы (CSS Profile)'],
   ['Лидерство: волонтёрство, общественная работа','Чаще считают автоматически по формам']]),
 'needblind':(['Need-blind','Need-aware'],[
   ['Финансы НЕ влияют на решение о приёме','Финансы учитываются при приёме'],
   ['После зачисления покрывают вашу нужду','Помощь есть, но бывает ограниченной'],
   ['Короткий список вузов (для иностранцев)','Большинство вузов США'],
   ['Нуждающимся студентам равные шансы','Нуждающимся сложнее поступить']]),
 'honors':(['Название (роль)','Описание (с цифрами и результатом)'],[
   ['National Olympiad Winner','11th year competing, 1 of 4 reps from school (>700 candidates), 1st prizes'],
   ['Writer, Screenwriter','Wrote literary works (200+) and on commission (10+), worked with 3 directors'],
   ['Researcher, Project Rep.','Ancient dramaturgy analysis with a VGIK professor, presented at conference, 2nd place'],
   ['(ваш пункт)','(коротко: роль и результат в цифрах)']]),
 'docs':(['Документ','Кто готовит','Откуда отправляется','Нотар. перевод'],[
   ['Рекомендательные письма','учителя школы','корпоративная почта школы','да'],
   ['Рекомендация работодателя','работодатель','почта компании','да'],
   ['Транскрипт оценок','школа','почта школы','да'],
   ['Аттестат','школа','почта школы','да'],
   ['Рекомендация профессора','профессор','почта профессора или вуза','да']]),
 'eaedrd':(['Тип подачи','Когда подаёшь','Когда ответ','Обязательства'],[
   ['Early Action','ноя-дек','фев-мар','не обязывает'],
   ['Early Decision','ноя-дек','фев-мар','обязывает, только 1 вуз'],
   ['Regular Decision','около 1 января','март-апрель','не обязывает']]),
}

doc.add_heading('Поступление в бакалавриат США', level=0)
sub=doc.add_paragraph('Пошаговый личный гайд Global Generation. Дорога от весны до подачи заявок: что делать, когда и в каком порядке.')
sub.runs[0].italic=True; sub.runs[0].font.color.rgb=RGBColor(0x55,0x55,0x55)

for t,v,pg in elements:
    if t=='h':
        doc.add_heading(clean(HDISP[v]), level=1)
    elif t=='img':
        p=doc.add_paragraph(clean(v)); shade(p,'EFF3F7')
        r=p.runs[0]; r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0x41,0x50,0x5f)
    elif t=='genpic':
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(v, width=Inches(5.9))
        cap=doc.add_paragraph('(иллюстрация, сгенерирована автоматически)'); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic=True; cap.runs[0].font.size=Pt(8); cap.runs[0].font.color.rgb=RGBColor(0x88,0x88,0x88)
    elif t=='table':
        make_table(*TABLES[v])
    elif t=='fn':
        p=doc.add_paragraph(); r=p.add_run(clean(v))
        r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x66,0x66,0x66)
        pPr=p._p.get_or_add_pPr(); pbd=OxmlElement('w:pBdr'); top=OxmlElement('w:top')
        for kk,vv in (('w:val','single'),('w:sz','4'),('w:space','2'),('w:color','CCCCCC')): top.set(qn(kk),vv)
        pbd.append(top); pPr.append(pbd)
    elif t=='pic':
        f,w,h=v
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.add_run()
        maxw,maxh=Inches(4.3),Inches(5.2); ratio=h/w if w else 1; width=maxw
        if width*ratio>maxh: width=Emu(int(maxh/ratio))
        run.add_picture(f, width=width)
        cap=doc.add_paragraph('(скриншот из оригинала)'); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic=True; cap.runs[0].font.size=Pt(8); cap.runs[0].font.color.rgb=RGBColor(0x88,0x88,0x88)
    else:
        txt=clean(v)
        CUES=('включают:','нужно собрать:','значат эти аббревиатуры:')
        if '●' in txt or '•' in txt:
            pos=min(txt.find(c) for c in '●•' if c in txt)
            intro=txt[:pos].strip()
            items=[x.strip() for x in re.split(r'[●•]', txt[pos:]) if x.strip()]
            if intro: doc.add_paragraph(intro)
            for it in items: doc.add_paragraph(it, style='List Bullet')
        elif any(c in txt for c in CUES) and ': - ' in txt:
            head, _, rest = txt.partition(': - ')
            doc.add_paragraph(head + ':')
            for it in re.split(r'\s-\s', rest):
                it=it.strip()
                if it: doc.add_paragraph(it, style='List Bullet')
        elif txt.startswith('- '):
            item=txt[2:].strip()
            doc.add_paragraph(item, style='List Bullet') if len(item)<=260 else doc.add_paragraph(item)
        else:
            doc.add_paragraph(txt)

OUT=BASE+"/Инструкция_Бакалавриат_ФИНАЛ.docx"
doc.save(OUT)
print("SAVED:",OUT)
print("footnotes:",sum(1 for t,_,_ in elements if t=='fn'),
      "| screenshots:",sum(1 for t,_,_ in elements if t=='pic'),
      "| KARTINKA:",sum(1 for t,_,_ in elements if t=='img'),
      "| elements:",len(elements))
